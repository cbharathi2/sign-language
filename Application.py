import streamlit as st
import cv2
import numpy as np
import mediapipe as mp
import tensorflow as tf
import time
from docx import Document
from io import BytesIO
from googletrans import Translator

# ---------------- PAGE CONFIG ----------------
st.set_page_config(
    page_title="AI Sign Language Translator Pro",
    layout="wide",
    initial_sidebar_state="expanded"
)

# ---------------- THEME ----------------
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=JetBrains+Mono:wght@400;600;700&family=Inter:wght@300;400;500;600&display=swap');

.main {background:#0a0a0a;color:#e0e0e0;}

.header-title{
font-family:'JetBrains Mono',monospace;
font-size:3rem;
font-weight:700;
text-align:center;
margin-bottom:2rem;
background:linear-gradient(135deg,#00d4ff,#7c3aed,#ff6b6b);
-webkit-background-clip:text;
-webkit-text-fill-color:transparent;
}

.stButton>button{
background:#1e1e2e;
color:white;
border-radius:8px;
border:1px solid #404040;
padding:0.7rem 1.5rem;
font-family:'JetBrains Mono';
}

.detection-container{
background:#1a1a2e;
border-radius:12px;
padding:2rem;
text-align:center;
border:1px solid #333;
}

.detection-text{
font-size:2.3rem;
font-family:'JetBrains Mono';
font-weight:700;
}

.output-container{
background:#1e1e2e;
padding:1.5rem;
border-radius:10px;
border-left:4px solid #00d4ff;
}

.output-text{
font-family:'JetBrains Mono';
font-size:1.2rem;
}

.confidence-meter{
height:8px;
background:#2a2a3a;
border-radius:4px;
}

.confidence-fill{
height:100%;
background:linear-gradient(90deg,#00d4ff,#7c3aed,#ff6b6b);
}

.status-label{
color:#a0a0a0;
font-size:0.8rem;
margin-bottom:5px;
font-family:Inter;
}
</style>
""", unsafe_allow_html=True)

# ---------------- TITLE ----------------
st.markdown('<h1 class="header-title"> Sign Language Translator Pro</h1>', unsafe_allow_html=True)

translator = Translator()

# ---------------- LOAD MODELS ----------------
@st.cache_resource
def load_models():

    asl_model = tf.keras.models.load_model("asl_model.h5")
    asl_classes = np.load("classes.npy", allow_pickle=True)

    hindi_model = tf.keras.models.load_model("hindimodal.h5")
    hindi_classes = np.load("hindhiclasses.npy", allow_pickle=True)

    word_model = tf.keras.models.load_model("acc.h5")
    word_classes = np.load("acc.npy", allow_pickle=True)

    return asl_model, asl_classes, hindi_model, hindi_classes, word_model, word_classes


asl_model, asl_classes, hindi_model, hindi_classes, word_model, word_classes = load_models()

# ---------------- SIDEBAR ----------------
st.sidebar.title("⚙ Settings")

mode = st.sidebar.selectbox(
    "Prediction Mode",
    [
        "ASL Letters",
        "Hindi Letters",
        "Words",
        "English → Tamil",
        "English → Hindi",
        "English → Malayalam"
    ]
)

st.session_state.hold_timer = st.sidebar.slider("Hold Time",1.0,5.0,2.0)

# ---------------- SESSION STATE ----------------
if "running" not in st.session_state:
    st.session_state.running = False

if "text" not in st.session_state:
    st.session_state.text = ""

if "current_detection" not in st.session_state:
    st.session_state.current_detection = ""

if "detection_start_time" not in st.session_state:
    st.session_state.detection_start_time = None

if "confidence" not in st.session_state:
    st.session_state.confidence = 0

# ---------------- MEDIAPIPE ----------------
mp_hands = mp.solutions.hands
mp_draw = mp.solutions.drawing_utils

hands = mp_hands.Hands(
    max_num_hands=2,
    min_detection_confidence=0.7,
    min_tracking_confidence=0.7
)

# ---------------- CONTROLS ----------------
col1, col2 = st.columns(2)

with col1:
    if st.button("▶ START CAMERA", use_container_width=True):
        st.session_state.running = True
        st.rerun()

with col2:
    if st.button("⏹ STOP CAMERA", use_container_width=True):
        st.session_state.running = False
        st.rerun()

# ---------------- LAYOUT ----------------
col_video, col_status = st.columns([2,1])

with col_video:
    video_placeholder = st.empty()

with col_status:
    detection_placeholder = st.empty()
    confidence_placeholder = st.empty()
    timer_placeholder = st.empty()

output_placeholder = st.empty()

# ---------------- TRANSLATION ----------------
def translate_word(word):

    try:
        if mode == "English → Tamil":
            return translator.translate(word, dest="ta").text
        if mode == "English → Hindi":
            return translator.translate(word, dest="hi").text
        if mode == "English → Malayalam":
            return translator.translate(word, dest="ml").text
        return word
    except:
        return word

# ---------------- CAMERA LOOP ----------------
if st.session_state.running:

    cap = cv2.VideoCapture(0)

    while st.session_state.running:

        ret,frame = cap.read()
        if not ret:
            break

        frame = cv2.flip(frame,1)
        rgb = cv2.cvtColor(frame, cv2.COLOR_BGR2RGB)

        result = hands.process(rgb)

        detected = ""
        confidence = 0
        features = []

        if result.multi_hand_landmarks:

            for hand in result.multi_hand_landmarks[:2]:

                mp_draw.draw_landmarks(frame,hand,mp_hands.HAND_CONNECTIONS)

                for lm in hand.landmark:
                    features.extend([lm.x,lm.y,lm.z])

        if len(features)>126:
            features=features[:126]

        while len(features)<126:
            features.append(0)

        features=np.array(features).reshape(1,126)

        # -------- WORD MODEL --------
        if mode == "Words" or mode.startswith("English →"):

            pred = word_model.predict(features,verbose=0)

            class_id=np.argmax(pred)
            confidence=float(np.max(pred))

            if confidence>0.80:

                word=word_classes[class_id]

                if mode.startswith("English →"):
                    word=translate_word(word)

                detected=word

        # -------- LETTER MODEL --------
        else:

            if result.multi_hand_landmarks:

                hand=result.multi_hand_landmarks[0]

                data=[]
                for lm in hand.landmark:
                    data.extend([lm.x,lm.y,lm.z])

                data=np.array(data).reshape(1,-1)

                if mode=="ASL Letters":

                    pred=asl_model.predict(data,verbose=0)
                    detected=asl_classes[np.argmax(pred)]

                if mode=="Hindi Letters":

                    pred=hindi_model.predict(data,verbose=0)
                    detected=hindi_classes[np.argmax(pred)]

        # -------- HOLD TIMER --------
        now=time.time()

        if detected==st.session_state.current_detection:

            if st.session_state.detection_start_time and now-st.session_state.detection_start_time>=st.session_state.hold_timer:

                st.session_state.text+=detected+" "
                st.session_state.current_detection=""
                st.session_state.detection_start_time=None

        else:

            st.session_state.current_detection=detected
            st.session_state.detection_start_time=now

        # -------- UI --------
        detection_placeholder.markdown(f"""
        <div class="detection-container">
        <div class="detection-text">{detected or "WAITING..."}</div>
        </div>
        """, unsafe_allow_html=True)

        conf_width=min(100,confidence*100)

        confidence_placeholder.markdown(f"""
        <div class="status-label">CONFIDENCE</div>
        <div style="color:#e0e0e0;font-weight:600">{confidence:.1%}</div>
        <div class="confidence-meter">
        <div class="confidence-fill" style="width:{conf_width}%"></div>
        </div>
        """, unsafe_allow_html=True)

        frame_rgb=cv2.cvtColor(frame,cv2.COLOR_BGR2RGB)
        video_placeholder.image(frame_rgb,use_column_width=True)

        output_placeholder.markdown(f"""
        <div class="output-container">
        <div class="output-text">{st.session_state.text or "No detections yet..."}</div>
        </div>
        """, unsafe_allow_html=True)

        time.sleep(0.03)

    cap.release()
    st.rerun()

# ---------------- DOWNLOAD ----------------
if st.session_state.text.strip():

    st.subheader("Export")

    doc=Document()
    doc.add_heading("Sign Language Output")
    doc.add_paragraph(st.session_state.text)

    buf=BytesIO()
    doc.save(buf)
    buf.seek(0)

    st.download_button("Download DOCX",buf,"translation.docx")

# ---------------- CLEAR ----------------
if st.button("Clear Text",use_container_width=True):

    st.session_state.text=""
    st.session_state.current_detection=""
    st.rerun()